#%%

"""
Packages for environment selection 
"""

import os  # Missing import for 'os'
from dotenv import find_dotenv, load_dotenv

import json
from schemas import SCHEMA_REGISTRY

"""
Packages for document writing
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
#from docx.shared import Inches

import openai

import pandas as pd
import re
import time
import pickle
import requests
from PyPDF2 import PdfReader
from io import BytesIO
from datetime import datetime
import configparser

import streamlit as st


def load_css(file_path):

    with open(file_path) as f:
        st.html(f'<style> {f.read()} </style>')



def get_pdf_files_in_directory(directory):
    """Returns a list of PDF files in the given directory."""
    return [file for file in os.listdir(directory) if file.endswith('.pdf')]




"""
==================================================================================================================
Assistant Creator and manager functions
==================================================================================================================
"""


def assistant_config(config, qualifier):

    res = {}

    model = config.get(f'assistant_{qualifier}', 'model', fallback=None)
    instructions = config.get(f'assistant_{qualifier}', 'instruction', fallback=None)
    temperature = config.getfloat(f'assistant_{qualifier}', 'temperature', fallback=None)
    topP = config.getfloat(f'assistant_{qualifier}', 'topP', fallback=None)

    res['model'] = model
    res['instructions'] = instructions
    res['temperature'] = temperature
    res['topP'] = topP
    return res 

def create_assistant(client, name, config):
    instructions = config['instructions']
    model = config['model']
    temp = config['temperature']
    topP = config['topP']

    assistant = client.beta.assistants.create(
        name=name,
        instructions=instructions,
        tools=[{"type": "file_search"}],
        model=model,
        temperature= temp,
        top_p= topP

    )
    return assistant.id  # Return the assistant ID

def load_file_to_assistant(client, vector_storeid ,
                           assistant_identifier, pdf_docs,
                           uploaded = True):

    # Get the current directory
    #current_directory = os.getcwd()

    # Get a list of PDF files in the current directory
    #pdf_files = get_pdf_files_in_directory(current_directory)

    #vector_store = client.beta.vector_stores.create(name="Business Overview")

    #pdf_dirs = [pdf._file_urls.upload_url for pdf in pdf_docs]
    
    #file_streams = [open(path, "rb") for path in pdf_files]
    #file_streams = [open(path, "rb") for path in pdf_dirs]

    if uploaded: 
        file_batch = client.beta.vector_stores.file_batches.upload_and_poll(
        vector_store_id= vector_storeid, files=pdf_docs
        )
        #st.write('\nUPLOADING THE DOCUMENTS\n')
        #st.write(file_batch.status)
        #st.write(file_batch.file_counts)
        #st.write('\n')
        #st.write(f'*'*20)
        #st.write(f'The file batches are:\n{file_batch}')
        #st.write(f'*'*20)
        #st.write('\n')
    else:

        # Open each file in binary mode
        file_streams = [open(file_path, "rb") for file_path in pdf_docs]

        try:
            # Upload the files to the vector store
            file_batch = client.beta.vector_stores.file_batches.upload_and_poll(
                vector_store_id= vector_storeid, files=file_streams
            )

            #st.write('\nUPLOADING THE DOCUMENTS\n')
            #st.write(file_batch.status)
            #st.write(file_batch.file_counts)


        finally:
            # Ensure all file streams are closed
            for stream in file_streams:
                stream.close()


    assistant = client.beta.assistants.update(
    assistant_id= assistant_identifier,
    tool_resources={"file_search": {"vector_store_ids": [vector_storeid]}},
    )



"""
==================================================================================================================
Assistant Question and Answering functions
==================================================================================================================
"""


import time

def get_answer(client, run, thread_identifier, timeout=120):
    start_time = time.time()
    while not run.status == "completed":
        if time.time() - start_time > timeout:
            raise TimeoutError("Assistant run timed out.")
        run = client.beta.threads.runs.retrieve(
            thread_id=thread_identifier,
            run_id=run.id
        )
    return run

"""
The following function is about loading the prompts we will use to fill the document.

This retrieves, from a .xlsx file, both the prompt and the placeholder metadata.

The placeholders corresponds to the ones in the .docx document and will be used 
select the appropriate place in which the assistant answer will be placed in 
the final document.
"""

def prompts_retriever(file_name, sheet_list):

    prompt_sheet = sheet_list[0]
    formatting_sheet = sheet_list[1]
        
    prompt_df = pd.read_excel(
        file_name, 
        sheet_name=prompt_sheet, 
        keep_default_na=True,  # Keep pandas default missing value recognition
        na_values=['']         # Treat empty strings as NaN
    )
    prompt_list = list(zip(prompt_df['Placeholder'], prompt_df['Prompt']))

    temp_df = pd.read_excel(file_name,sheet_name=formatting_sheet)

    additional_formatting_requirements = temp_df.iloc[0,0]

    return prompt_list, additional_formatting_requirements, prompt_df


def prompt_creator(prompt_df, prompt_name, 
                   prompt_message, additional_formatting_requirements,
                   answers_dict, json_dict = {}, item = ''):
    
    print('&'*40) 
    print(prompt_message)

    row = prompt_df[prompt_df['Placeholder'] == prompt_name]

    if pd.isna(row['Links'].iloc[0]):

        prompt_message_format = prompt_message + additional_formatting_requirements

    else:

        reference = answers_dict[row['Links'].iloc[0]]
        prompt_message_format = reference
        prompt_message_format += prompt_message + additional_formatting_requirements

    if pd.isna(row['JSON_use'].iloc[0]):
        pass

    else: 
        json_reference = row['JSON_use'].iloc[0]
        #st.write(f'{json_reference}')
        json_output = json_dict[json_reference]
        #st.write(f'{json_output}')
        json_output_str = '\n'.join(json_output)
        prompt_message_format = f'{prompt_message_format}'
        prompt_message_format = prompt_message_format.replace("{json_output}", json_output_str)
        #prompt_message_format = prompt_message_format.format(json_output = json_output)
    try:
        if pd.isna(row['For_loop'].iloc[0]):

            pass

        else: 

            prompt_message_format = f'{prompt_message_format}'
            prompt_message_format = prompt_message_format.replace("{sector}", item)

    except:
        pass

    
    """
    We will iterate through all the prompts that are present in the .xlsx file.

    The prompt_list object is a list of tuples with:

        1) prompt_name = The placeholder in the .docx file that is associated with the 
        current prompt.

        2) prompt_message = The prompt itself that will be used to ask the assistant a question.
    """
    
    print(prompt_message_format)
    
    return prompt_message_format
    



def separate_thread_answers(client, prompt_message_format,
                            assistant_identifier, 
                            same_chat = False, thread_id = ''):
    if not same_chat: 

        thread = client.beta.threads.create()
        thread_identifier = thread.id

    else:

        thread_identifier = thread_id

    """
    We essentially append our message to the current thread, to query the assistant
    """

    user_message = client.beta.threads.messages.create(
        thread_id=thread_identifier,
        role="user",
        content=prompt_message_format
    )

    """
    This is the actual interaction with the OpenAI assistant
    """

    run = client.beta.threads.runs.create(
        thread_id=thread_identifier,
        assistant_id= assistant_identifier  
    )

    """
    In order to achieve a sequential workflow in which we move to the next prompt 
    only when the previous one was answered, we added this while loop to prevent
    moving forward until prompt completion.

    """
    
    run = get_answer(client, run, thread_identifier)
  
    
    """
    We retrieve the entire list of messages that are part of the thread.

    By looping through the data attribute we are moving from the last message 
    upwards to the first.

!!!!!!!!!!!!!!!!!!!
    We will retrieve the first message that the answer from the assistant
    whose content is textual.
!!!!!!!!!!!!!!!!!!!
    """

    messages = client.beta.threads.messages.list(thread_id=thread_identifier)
    assistant_response = None

    for message in messages.data:  
        if message.role == "assistant":
            for content_block in message.content:
                if content_block.type == "text":
                    assistant_response = content_block.text.value
                    break
            if assistant_response:
                break
    print(assistant_response)
    return assistant_response, thread_identifier


def json_schema_answer(client, prompt_df, prompt_name, json_dict,
                       assistant_response):

    row = prompt_df[prompt_df['Placeholder'] == prompt_name]

    if pd.isna(row['JSON_make'].iloc[0]):

        return json_dict

    else:
        #st.write(f'entered in the correct control flow')
        schema_name = row['JSON_make'].iloc[0]

        json_schema = SCHEMA_REGISTRY[schema_name]

        prompt = assistant_response

        question = row['JSON_question'].iloc[0]

        prompt += question

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": "You are a Private Equity Analyst assistant"},
            {"role": "user", "content": prompt}],
            
            response_format={
                "type": "json_schema",
                "json_schema": json_schema
            }
        )
        #st.write(f'{response}')
        json_string = response.choices[0].message.content

        #st.write(f'{json_string}')

        # Parse the JSON string into a Python dictionary
        parsed_json = json.loads(json_string)

        #st.write(f'{parsed_json}')

        # Extract the Python list
        companies_list = parsed_json.get("companies", [])

        #st.write(f'{companies_list}')
        json_dict[schema_name] = companies_list

        return json_dict


def missing_warning(client, thread_id, prompt, assistant_identifier):

    question = """Given that, with the information in the files uploaded to the assistant, the model was not able to answer the following question:\n"""
    question = question + prompt

    question = """
    Please write what will be a warning to the user that the model was not able to find the answer.

    It should follow: "The AI Assistant did not find/was not confident enough to write about: {the theme of the question}.
    Please try to be concise, as this is meant to function as a warning to the user, not as a full-fledged answer to a prompt please.
    """

    warning, x = separate_thread_answers(client, prompt, assistant_identifier)
    #warning = "Highlight!$% " + warning
    #warning += " Highlight!$%"

    return warning

def warning_check(answer, client, thread_id, prompt, assistant_identifier):

    if "not_found" not in answer.lower():

        highlight = False

        return answer, highlight
    
    else:

        warning = missing_warning(client, thread_id, prompt, assistant_identifier)
        highlight = True

        #st.write(f'To the prompt: {prompt}')
        #st.write(f'Gives waring: {warning}')
        #st.write(f'The highlight parameter is: {highlight}')

        

        return warning, highlight


"""
==================================================================================================================
REFERENCE MARKET PART
==================================================================================================================
"""

def get_pdf_text(pdf_docs): 
    """Extracts text from a list of PDF files."""
    text = "" 
    for pdf in pdf_docs: 
        pdf_reader = PdfReader(pdf) 

        for page in pdf_reader.pages: 
            text += page.extract_text() or "" 
                
    return text

def html_retriever(uploaded_files):
    #st.write(f"{uploaded_files}")

    html_dir = "retrieved_html_files"
    os.makedirs(html_dir, exist_ok=True)

    extracted_text = get_pdf_text(uploaded_files)
    #st.write(f"{extracted_text}")

    """
    Finding the URLs inside the files
    """
        
    url_df = pd.DataFrame(columns = ['single_line', 'double_line', 'triple_line'])
    url_pattern = r'https?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    lines = extracted_text.split("\n")

    found = []

    for idx, line in enumerate(lines):

        if re.search(r"https?://", line):

            l = []
            s = re.findall(url_pattern, line)
            l.append(s[0])

            l.append(s[0] + lines[idx + 1])
            l.append(s[0] + "".join(lines[idx+1:idx+3]))

            url_df.loc[len(url_df)] = l

    
        
    # Ensure the output directory exists
    #os.makedirs('html_files', exist_ok=True)

    # List to hold file paths
    html_file_paths = []

    # Download HTML content for each URL
    for idx in url_df.index:
        for column in url_df.columns:
            url = url_df.loc[idx, column]
            try:
                response = requests.get(url)
                response.raise_for_status()  # Check for HTTP errors

                # Save HTML content to a file
                file_name = f"page_{idx}.html"
                file_path = os.path.join(html_dir, file_name)
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(response.text)

                html_file_paths.append(file_path)
                #st.write(f"HTML content retrieved and saved from {url}")
                break  # Stop trying other columns once successful
            except Exception as e:
                #st.warning(f"Failed to fetch {url}: {e}")
                continue

    return html_file_paths

    
"""
==================================================================================================================
App design functions
==================================================================================================================
"""

def update_progressbar(progress_bar, message_placeholder,
                       milestone, steps,
                       message):
    progress_bar.progress(milestone[0] / steps)
    message_placeholder.markdown(message)
    time.sleep(1)  
    milestone[0] += 1


"""
==================================================================================================================
DOCUMENT FORMATTING
==================================================================================================================
"""

def format_spec(config):
    font_size = config.get('document_format', 'font_size', fallback=None)
    font_type = config.get('document_format', 'font_type', fallback=None)

    return font_size, font_type

def remove_source_patterns(text):
    """
    Removes patterns like   from the     
    Args:
        text (str): The input string containing potential patterns to remove.
    
    Returns:
        str: The cleaned text without the specified patterns.
    """
    # Define the regular expression to match the pattern
    pattern = r"【\d+:\d+†source】"
    
    # Use re.sub to remove all occurrences of the pattern
    cleaned_text = re.sub(pattern, "", text)
    
    # Return the cleaned text
    return cleaned_text


def document_filler(doc_copy, prompt_name, assistant_response, last_p, 
                    section_creator = False, highlighting=False):
    #First we loop through all the paragraphs.
    last_par_idx = last_p
    if not section_creator:
        for i, paragraph in enumerate(doc_copy.paragraphs):

            #If the prompt_name correspond to the placeholder making up the paragraph
            #we move to the filling part
            if prompt_name in paragraph.text:
                
                #This is for formatting reasons to avoid alignment problems
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                #Then, we move to the run objects inside the paragraph.
                #The reason is that in this way, when we replace the placeholder 
                #we will keep the placeholder's formatting
                for run in paragraph.runs:
                    if prompt_name in run.text:
                        run.text = run.text.replace(prompt_name, assistant_response)
                        if highlighting:
                            #st.write(f'CURRENTLY HIGHLIGHTING: {prompt_name}')
                            # Create a shading XML element to highlight the run
                            shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF00" w:val="clear"/>'.format(nsdecls('w')))
                            run._r.get_or_add_rPr().append(shading_elm)

                        last_par_idx = i

    else: 
        config = configparser.ConfigParser()
    
        # Read the .cfg file
        config.read('assistant_config.cfg') 

        font_size = config.get('document_format', 'font_size', fallback=None)
        font_size = int(font_size)
        font_type = config.get('document_format', 'font_type', fallback=None)

        new_section_par = insert_paragraph_after(doc_copy.paragraphs[last_p],
                                                 text= assistant_response,
                                                 font_size= font_size, 
                                                 font_type= font_type)
        # After insertion, refresh the index by finding new_section_par in doc_copy.paragraphs
        # The newly inserted paragraph will appear right after last_modified_idx.
        # Since we inserted one after it, it should now be at last_modified_idx + 1.

        if highlighting:
            #st.write(f'The {prompt_name} is being highlighted')
            # Apply shading to the entire paragraph or each run if needed
            for run in new_section_par.runs:
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF00" w:val="clear"/>'.format(nsdecls('w')))
                run._r.get_or_add_rPr().append(shading_elm)
        
        # One approach is to directly compute the index:
        new_section_idx = last_p + 1

        # Update last_modified_idx to this new paragraph
        last_par_idx = new_section_idx

    return last_par_idx
        



def insert_paragraph_after(paragraph, text=None, style=None, section = False, 
                           font_size = None, font_type = None, last_p = None, doc_copy = None):

    if not section:
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        new_paragraph = paragraph._parent.add_paragraph()
        new_paragraph._p = new_p
        #new_paragraph.style = 'List Bullet'
        if text:
            run = new_paragraph.add_run(text)
            run.bold = False
            run.font.bold = False
        
        # Explicitly reset italic and underline
            run.italic = False
            run.font.italic = False
            run.underline = None
            run.font.underline = False

        return new_paragraph

    else:

        last_paragraph = paragraph._parent.paragraphs[last_p]
        new_p = OxmlElement('w:p')
        last_paragraph._p.addnext(new_p)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        new_paragraph = paragraph._parent.add_paragraph()
        new_paragraph._p = new_p
        #new_paragraph.style = 'List Bullet'
        if text:
            run = new_paragraph.add_run(text)
            # Set the font properties
            run.font.size = Pt(font_size)
            run.font.name = font_type
            run.bold = True
        
        last_p += 1
        
        return new_paragraph, last_p
    

from markdownify import markdownify as md

def document_filler_2(doc_copy, prompt_name, assistant_response, last_p, 
                    section_creator=False, highlighting=False):
    """
    Function to replace placeholders in a .docx document with responses.
    Includes support for Markdown formatting and highlighting.
    
    Parameters:
        doc_copy (Document): A `python-docx` Document object.
        prompt_name (str): Placeholder name to replace in the document.
        assistant_response (str): Text to insert in place of the placeholder.
        last_p (int): Index of the last modified paragraph.
        section_creator (bool): Whether to create a new section for the response.
        highlighting (bool): Whether to apply highlighting to the inserted text.
    Returns:
        int: Index of the last modified paragraph.
    """
    # Normalize assistant_response by converting Markdown to plain text or styled content

    markdown_chars = [
    '**','#', '##', '###', '####', '#####', '######'   # Section symbols
    ]
    if any(char in assistant_response for char in markdown_chars):

        assistant_response = md(assistant_response)

    # Last modified paragraph index
    last_par_idx = last_p

    if not section_creator:
        for i, paragraph in enumerate(doc_copy.paragraphs):

            # Check if the placeholder is in the paragraph text
            if prompt_name in paragraph.text:

                # Align paragraph to the left
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Replace placeholder text
                full_text = paragraph.text.replace(prompt_name, assistant_response)

                # Clear existing runs
                for run in paragraph.runs:
                    run.clear()
                
                # Create a new run with the updated text
                new_run = paragraph.add_run(full_text)

                new_run.bold = False
                new_run.italic = False
                new_run.underline = None
                # Apply highlighting if required
                if highlighting:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF00" w:val="clear"/>'.format(nsdecls('w')))
                    new_run._r.get_or_add_rPr().append(shading_elm)

                last_par_idx = i
                break  # Stop after replacing the first instance of the placeholder

    else:
        # Handle section creation logic
        config = configparser.ConfigParser()
        config.read('assistant_config.cfg')

        # Get font size and type from config
        font_size = int(config.get('document_format', 'font_size', fallback=12))
        font_type = config.get('document_format', 'font_type', fallback='Times New Roman')
        st.write(f'The prompt is: {prompt_name} and pargraph: {last_p}')
        # Insert new paragraph after the last modified one
        new_section_par, last_par_idx = insert_paragraph_after(doc_copy.paragraphs[last_p],
                                                               text=assistant_response,
                                                               section=True,
                                                               font_size=font_size,
                                                               font_type=font_type,
                                                               last_p=last_p,
                                                               doc_copy=doc_copy)

        # Highlight new paragraph if needed
        if highlighting:
            for run in new_section_par.runs:
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF00" w:val="clear"/>'.format(nsdecls('w')))
                run._r.get_or_add_rPr().append(shading_elm)

    return last_par_idx


def adding_headers(document, title):

    section = document.sections[0]

    # Access the header of the section
    header = section.header

    paragraph = header.paragraphs[0]

    new_text = []

    left_paragraph= f"Project {title.capitalize()}"
    
    new_text.append(left_paragraph)

    current_date = datetime.now()

    right_paragraph = current_date.strftime("%B %Y").capitalize()
    #print(right_paragraph)

    new_text.append(right_paragraph)

    #paragraph.text = f"{left_paragraph}\t\t{right_paragraph}"
    
    l = ['Project', 'Date']

    for sub, new in zip(l, new_text):

        for run in paragraph.runs:
                    print(run.text)
                    if sub in run.text:
                        run.text = run.text.replace(sub, new)


from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

#def highlight_paragraphs_with_keyword(doc_path, keyword):

def highlight_paragraphs_with_keyword(doc, keyword, font_name, font_size):

    print(f'The font name is: {font_name}')
    print(f'The font size is: {font_size}')
    # Iterate through each paragraph
    for paragraph in doc.paragraphs:
        if keyword in paragraph.text:
            # Highlight the paragraph
            shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
            paragraph._p.get_or_add_pPr().append(shading_elm)
            
            # Remove the keyword
            paragraph.text = paragraph.text.replace(keyword, '').strip()
            
            # Apply font and size to all runs in the paragraph
            for run in paragraph.runs:
                run.font.name = font_name  # Set the font name
                run.font.size = Pt(font_size)  # Set the font size

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt    

def highlight_sections(doc, keyword, font_name, font_size):
    print(f'The font name is: {font_name}')
    print(f'The font size is: {font_size}')
    
    highlighting = False  # Indicates if we are currently in a highlight section
    for paragraph in doc.paragraphs:
        text = paragraph.text

        if keyword in text: 

            paragraph.text = paragraph.text.replace(keyword, '')
            highlighting = not highlighting

        for run in paragraph.runs:
            text = run.text
            # Check if there's a highlight marker in this run
            if keyword in text:
                # Remove the keyword from the run text
                run.text = run.text.replace(keyword, '')
                # Toggle the highlighting state
                highlighting = not highlighting
                # After toggling, if highlighting is now on, highlight the run
                # If highlighting just turned off, this run will end highlighting.
            
            # If currently highlighting, apply highlight and formatting
            if highlighting:
                # Add shading to this run
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                rPr = run._r.get_or_add_rPr()
                rPr.append(shading_elm)

                # Apply font and size
                run.font.name = font_name
                run.font.size = Pt(font_size)



from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def boldify_text_between_asterisks(doc_path):
    """
    Processes a Word document to replace text surrounded by asterisks with bold text.
    
    Args:
        doc_path (str): Path to the input .docx file.
        output_path (str): Path to save the modified .docx file.
    """
    # Open the document

    # Iterate through all paragraphs in the document
    for paragraph in doc_path.paragraphs:
        # Split the paragraph text into runs for processing
        for run in paragraph.runs:
            # If the text contains asterisks
            if "\*\*" in run.text:
                parts = run.text.split("\*\*")
                new_text = []
                
                # Iterate over parts to process asterisk-marked text
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Text between asterisks
                        bold_run = paragraph.add_run(part)
                        bold_run.bold = True
                        # Copy formatting from the current run
                        bold_run.font.size = run.font.size
                        bold_run.font.name = run.font.name
                        bold_run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
                    else:  # Regular text
                        normal_run = paragraph.add_run(part)
                        # Copy formatting from the current run
                        normal_run.font.size = run.font.size
                        normal_run.font.name = run.font.name
                        normal_run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
                
                # Clear the original run text
                run.clear()
